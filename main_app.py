"""
EDA desde Texto - Extracción de tablas con LLM (Groq / Llama 3.3 70B)
Incluye dos modos:
  1) Extracción general de tabla desde un párrafo con cifras + EDA.
  2) Evaluación de entrevistas desde texto -> tabla de puntajes -> radar interactivo.
Ejecutar con: streamlit run main_app.py
"""

import io
import json
import re

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document as DocxDocument
from groq import Groq, APIError, AuthenticationError, RateLimitError
from scipy import stats

# --------------------------------------------------------------------------------------
# CONFIGURACIÓN GENERAL
# --------------------------------------------------------------------------------------
st.set_page_config(
    page_title="EDA desde Texto (LLM)",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded",
)

MODEL_ID = "llama-3.3-70b-versatile"

EXTRACTION_SYSTEM_PROMPT = """Eres un asistente experto en extracción de datos estructurados a partir de texto en español.

Se te dará un párrafo que contiene cifras (números, porcentajes, montos, cantidades, fechas, etc.)
asociadas a entidades, categorías o periodos. Tu tarea es:

1. Identificar cada cifra relevante y la entidad/categoría/periodo al que pertenece.
2. Construir una tabla en forma de arreglo JSON de objetos, donde cada objeto es una fila.
3. Usar las MISMAS claves (nombres de columna) en todos los objetos, de forma consistente.
4. Los valores numéricos deben quedar como números (int o float), no como texto ni con símbolos (%, $, comas de miles).
5. Si el texto no permite construir una tabla clara, extrae lo que puedas de forma razonable.

Responde ÚNICAMENTE con el arreglo JSON, sin explicaciones, sin texto adicional, sin backticks
de markdown y sin comentarios. Ejemplo de formato de salida:

[
  {"Entidad": "Antioquia", "Variable": "Casos reportados", "Valor": 1450, "Periodo": "2024"},
  {"Entidad": "Valle del Cauca", "Variable": "Casos reportados", "Valor": 980, "Periodo": "2024"}
]
"""

EXAMPLE_TEXT_GENERAL = (
    "Durante el primer trimestre de 2024, el departamento de Antioquia reportó 1450 "
    "casos de dengue, un aumento del 23% frente al mismo periodo de 2023. En Valle del "
    "Cauca se registraron 980 casos, con una letalidad del 0.4%. Cundinamarca, por su "
    "parte, presentó 620 casos y una cobertura de fumigación del 65%. A nivel nacional, "
    "el presupuesto asignado para el control vectorial fue de 12500 millones de pesos, "
    "distribuidos en 32 departamentos."
)

EXAMPLE_TEXT_ENTREVISTA = """Candidato: Mariana Restrepo Gómez | Documento: CC 1.037.612.489
Pregunta: Cuéntanos tu proceso creativo para diseñar un personaje.
Respuesta: Empiezo con una hoja de personalidad y un moodboard, luego exploro siluetas y por
último detallo vestuario que refuerce la historia del personaje. Colaboro estrechamente con
narrativa desde el inicio y tomo muy bien el feedback de dirección de arte, aunque mi manejo
de herramientas 3D es intermedio, no avanzado.

Candidato: Andrés Felipe Cárdenas Ríos | Documento: CC 71.345.612
Pregunta: Cuéntanos tu proceso creativo para diseñar un personaje.
Respuesta: Domino ZBrush, Maya y Substance Painter a nivel experto y calculo el presupuesto de
polígonos desde el día uno. La narrativa no es mi fuerte, dependo del equipo de concept para
el trasfondo. Colaboro bien con animación y programación, pero menos con narrativa.
"""

DEFAULT_AXES = "Creatividad, Dominio técnico, Narrativa, Colaboración, Adaptabilidad, Profesionalismo"


# --------------------------------------------------------------------------------------
# SIDEBAR: API KEY Y CONFIGURACIÓN GENERAL
# --------------------------------------------------------------------------------------
with st.sidebar:
    st.header("⚙️ Configuración")
    api_key = st.text_input(
        "API Key de Groq",
        type="password",
        placeholder="gsk_...",
        help="No se guarda en ningún lado; solo se usa durante esta sesión.",
    )
    st.caption(f"Modelo: `{MODEL_ID}`")

    with st.expander("Parámetros del modelo"):
        temperature = st.slider("Temperatura", 0.0, 1.0, 0.1, 0.05,
                                 help="Valores bajos = extracción/evaluación más consistente.")
        max_tokens = st.slider("Máx. tokens de respuesta", 256, 4096, 2000, 128)

    st.divider()
    modo = st.radio(
        "Modo de análisis",
        ["📊 Tabla general desde texto", "🕸️ Evaluación de entrevista (radar)"],
        index=0,
    )

    st.divider()
    st.caption("Obtén tu API Key gratuita en [console.groq.com/keys](https://console.groq.com/keys).")

st.title("📝 EDA desde Texto — Extracción con LLM")

if not api_key:
    st.info("⬅️ Ingresa tu **API Key de Groq** en la barra lateral para comenzar.")
    st.stop()

client = Groq(api_key=api_key)


# --------------------------------------------------------------------------------------
# UTILIDADES COMPARTIDAS
# --------------------------------------------------------------------------------------
def clean_json_response(raw_text: str) -> str:
    """Quita posibles fences de markdown y texto sobrante alrededor del JSON."""
    text = raw_text.strip()
    text = re.sub(r"^```(json)?", "", text.strip(), flags=re.IGNORECASE).strip()
    text = re.sub(r"```$", "", text.strip()).strip()
    start = text.find("[")
    end = text.rfind("]")
    if start != -1 and end != -1 and end > start:
        text = text[start:end + 1]
    return text


def call_llm_for_json(system_prompt: str, user_text: str):
    response = client.chat.completions.create(
        model=MODEL_ID,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_text},
        ],
        temperature=temperature,
        max_tokens=max_tokens,
    )
    raw = response.choices[0].message.content
    cleaned = clean_json_response(raw)
    data = json.loads(cleaned)
    if isinstance(data, dict):
        data = [data]
    return pd.DataFrame(data), raw


def handle_llm_errors(fn, *args, **kwargs):
    """Ejecuta fn con manejo estándar de errores de Groq / JSON. Devuelve (df, raw) o (None, None)."""
    try:
        return fn(*args, **kwargs)
    except AuthenticationError:
        st.error("❌ API Key inválida. Verifica que la copiaste correctamente desde Groq Console.")
    except RateLimitError:
        st.error("⏳ Se alcanzó el límite de solicitudes de Groq. Espera unos segundos e intenta de nuevo.")
    except json.JSONDecodeError:
        st.error(
            "⚠️ El modelo no devolvió un JSON válido. Intenta de nuevo, baja la temperatura "
            "o reformula el texto de entrada."
        )
    except APIError as e:
        st.error(f"⚠️ Error de la API de Groq: {e}")
    except Exception as e:
        st.error(f"⚠️ Ocurrió un error inesperado: {e}")
    return None, None


# ========================================================================================
# MODO 1: TABLA GENERAL DESDE TEXTO + EDA
# ========================================================================================
if modo == "📊 Tabla general desde texto":
    st.caption(
        "Pega un párrafo con cifras (informes, noticias, boletines) y el modelo lo convierte "
        "en una tabla estructurada para analizarla."
    )

    if "input_text" not in st.session_state:
        st.session_state.input_text = ""

    c1, c2 = st.columns([4, 1])
    with c1:
        st.text_area(
            "Pega aquí tu párrafo con cifras",
            key="input_text",
            height=180,
            placeholder="Ejemplo: 'Durante 2024, la empresa X reportó ventas de 3200 millones...'",
        )
    with c2:
        st.write("")
        st.write("")
        if st.button("📋 Usar texto de ejemplo", use_container_width=True, key="ejemplo_general"):
            st.session_state.input_text = EXAMPLE_TEXT_GENERAL
            st.rerun()

    extract_clicked = st.button("🔍 Extraer tabla con el LLM", type="primary", use_container_width=True)

    if extract_clicked:
        if not st.session_state.input_text.strip():
            st.warning("Pega un párrafo con cifras antes de extraer la tabla.")
        else:
            with st.spinner("Extrayendo tabla con el modelo..."):
                table_df, raw_response = handle_llm_errors(
                    call_llm_for_json, EXTRACTION_SYSTEM_PROMPT, st.session_state.input_text
                )
                if table_df is not None:
                    st.session_state.extracted_df = table_df
                    st.session_state.raw_response = raw_response

    if "extracted_df" not in st.session_state:
        st.stop()

    st.divider()
    st.subheader("📊 Tabla extraída (editable)")
    st.caption("Puedes corregir manualmente cualquier celda antes de continuar con el EDA.")

    edited_df = st.data_editor(
        st.session_state.extracted_df, num_rows="dynamic", use_container_width=True, key="editor_general"
    )

    with st.expander("Ver JSON crudo devuelto por el modelo"):
        st.code(st.session_state.get("raw_response", ""), language="json")

    csv_buf = io.StringIO()
    edited_df.to_csv(csv_buf, index=False)
    st.download_button("⬇️ Descargar tabla como CSV", csv_buf.getvalue(),
                        file_name="tabla_extraida.csv", mime="text/csv")

    if edited_df.empty:
        st.warning("La tabla está vacía. Ajusta el texto de entrada o edita la tabla manualmente.")
        st.stop()

    df = edited_df.copy()
    numeric_cols, categorical_cols = [], []
    for col in df.columns:
        coerced = pd.to_numeric(df[col], errors="coerce")
        if coerced.notna().sum() >= max(1, int(len(df) * 0.6)):
            df[col] = coerced
            numeric_cols.append(col)
        else:
            categorical_cols.append(col)

    st.divider()
    st.subheader("🔎 Análisis exploratorio de la tabla extraída")

    tab_quant, tab_qual, tab_graph = st.tabs(["🔢 Cuantitativo", "🔤 Cualitativo", "📈 Gráfico"])

    with tab_quant:
        if numeric_cols:
            desc = df[numeric_cols].describe().T
            desc["mediana"] = df[numeric_cols].median()
            st.dataframe(desc.style.format(precision=2), use_container_width=True)

            if len(df) >= 3:
                col_sel = st.selectbox("Variable para prueba de normalidad (Shapiro-Wilk)", numeric_cols)
                sample = df[col_sel].dropna()
                if len(sample) >= 3:
                    stat, p_value = stats.shapiro(sample)
                    n1, n2 = st.columns(2)
                    n1.metric("Estadístico W", f"{stat:.4f}")
                    n2.metric("p-valor", f"{p_value:.4f}")
        else:
            st.info("No se detectaron columnas numéricas en la tabla extraída.")

    with tab_qual:
        if categorical_cols:
            cat_sel = st.selectbox("Variable categórica", categorical_cols)
            freq = df[cat_sel].value_counts(dropna=False).rename("Frecuencia").to_frame()
            c1, c2 = st.columns([1, 1.4])
            with c1:
                st.dataframe(freq, use_container_width=True)
            with c2:
                fig = px.bar(freq.reset_index(), x=cat_sel, y="Frecuencia", text="Frecuencia",
                             title=f"Distribución de {cat_sel}")
                st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("No se detectaron columnas categóricas en la tabla extraída.")

    with tab_graph:
        if numeric_cols and categorical_cols:
            c1, c2 = st.columns(2)
            with c1:
                cat_x = st.selectbox("Variable categórica (eje X)", categorical_cols, key="g_cat")
            with c2:
                num_y = st.selectbox("Variable numérica (eje Y)", numeric_cols, key="g_num")
            fig_bar = px.bar(df, x=cat_x, y=num_y, color=cat_x, title=f"{num_y} por {cat_x}")
            st.plotly_chart(fig_bar, use_container_width=True)

        if len(numeric_cols) >= 2:
            c3, c4 = st.columns(2)
            with c3:
                x_var = st.selectbox("Eje X", numeric_cols, index=0, key="s_x")
            with c4:
                y_var = st.selectbox("Eje Y", numeric_cols, index=min(1, len(numeric_cols) - 1), key="s_y")
            fig_scatter = px.scatter(df, x=x_var, y=y_var, title=f"{y_var} vs {x_var}")
            st.plotly_chart(fig_scatter, use_container_width=True)

            st.markdown("**Matriz de correlación**")
            corr = df[numeric_cols].corr()
            fig_corr = px.imshow(corr, text_auto=".2f", color_continuous_scale="RdBu_r", zmin=-1, zmax=1)
            st.plotly_chart(fig_corr, use_container_width=True)

        if not numeric_cols and not categorical_cols:
            st.info("No hay suficientes columnas para generar gráficos.")


# ========================================================================================
# MODO 2: EVALUACIÓN DE ENTREVISTA DESDE TEXTO -> TABLA DE PUNTAJES -> RADAR INTERACTIVO
# ========================================================================================
else:
    st.caption(
        "Pega el texto de una o varias entrevistas (nombre, documento y respuestas de cada "
        "candidato). El LLM evalúa cualitativamente cada respuesta y genera puntajes por "
        "competencia; luego puedes comparar a los candidatos con un radar interactivo."
    )

    with st.sidebar:
        st.divider()
        st.subheader("🕸️ Ejes de evaluación")
        axes_input = st.text_area(
            "Competencias a evaluar (separadas por coma)",
            value=DEFAULT_AXES,
            height=80,
            help="El LLM asignará un puntaje de 1 a `escala_max` en cada uno de estos ejes.",
        )
        scale_max = st.number_input("Escala máxima de puntaje", min_value=5, max_value=100, value=10, step=5)

    axes_list = [a.strip() for a in axes_input.split(",") if a.strip()]

    if "interview_text" not in st.session_state:
        st.session_state.interview_text = ""

    st.markdown("**📄 Carga el documento con todas las entrevistas**")
    uploaded_interview_file = st.file_uploader(
        "Sube el archivo con las entrevistas (.docx o .txt)",
        type=["docx", "txt"],
        key="interview_file",
        help="Por ejemplo, el Word consolidado con nombre, documento y respuestas de cada candidato.",
    )

    def extract_text_from_docx(file_obj) -> str:
        doc = DocxDocument(file_obj)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells_text = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells_text:
                    parts.append(" | ".join(cells_text))
        return "\n".join(parts)

    if uploaded_interview_file is not None and st.session_state.get("last_uploaded_name") != uploaded_interview_file.name:
        try:
            if uploaded_interview_file.name.lower().endswith(".docx"):
                extracted_text = extract_text_from_docx(uploaded_interview_file)
            else:
                extracted_text = uploaded_interview_file.read().decode("utf-8", errors="ignore")
            st.session_state.interview_text = extracted_text
            st.session_state.last_uploaded_name = uploaded_interview_file.name
            st.rerun()
        except Exception as e:
            st.error(f"⚠️ No se pudo leer el archivo: {e}")

    with st.expander("✏️ Ver / editar el texto cargado manualmente", expanded=not st.session_state.interview_text):
        st.text_area(
            "Texto de las entrevistas (editable)",
            key="interview_text",
            height=260,
            placeholder="Candidato: Nombre | Documento: CC 0000000\nPregunta: ...\nRespuesta: ...",
        )
        if st.button("📋 Usar texto de ejemplo", key="ejemplo_entrevista"):
            st.session_state.interview_text = EXAMPLE_TEXT_ENTREVISTA
            st.rerun()

    n_chars = len(st.session_state.interview_text)
    if n_chars > 0:
        st.caption(f"Texto cargado: {n_chars:,} caracteres (~{n_chars // 4:,} tokens aprox.).")
    if n_chars > 60000:
        st.warning(
            "El documento es bastante largo; si el modelo trunca la respuesta o falla, "
            "sube el archivo por partes o aumenta 'Máx. tokens de respuesta' en la barra lateral."
        )

    evaluate_clicked = st.button("🕸️ Evaluar entrevista(s) y generar puntajes", type="primary", use_container_width=True)

    def build_interview_prompt(axes, scale):
        axes_str = ", ".join(f'"{a}"' for a in axes)
        example_obj = {"Candidato": "Nombre Ejemplo", "Documento": "CC 0000000"}
        for a in axes:
            example_obj[a] = round(scale * 0.7)
        return f"""Eres un evaluador experto de entrevistas de selección de personal, especializado en
roles creativos y técnicos de la industria de videojuegos.

Se te dará el texto de una o varias entrevistas. Cada entrevista puede incluir el nombre del
candidato, su número de documento, y una o varias preguntas con sus respuestas.

Tu tarea:
1. Identifica cada candidato distinto mencionado en el texto (por nombre y/o documento).
2. Evalúa cualitativamente sus respuestas y asigna un puntaje entero de 1 a {scale} en cada uno
   de estos ejes de competencia: {axes_str}.
3. Si el texto no da información suficiente para evaluar un eje de un candidato, asigna tu mejor
   estimación razonable en vez de dejarlo vacío.
4. Devuelve un arreglo JSON con un objeto por candidato, usando EXACTAMENTE estas claves:
   "Candidato", "Documento" y una clave por cada eje de competencia (con el mismo nombre y
   mayúsculas/tildes que se te dieron).

Responde ÚNICAMENTE con el arreglo JSON, sin explicaciones ni texto adicional ni backticks de
markdown. Ejemplo de formato de salida (con un solo candidato):

[{json.dumps(example_obj, ensure_ascii=False)}]
"""

    if evaluate_clicked:
        if not st.session_state.interview_text.strip():
            st.warning("Pega el texto de al menos una entrevista antes de evaluar.")
        elif not axes_list:
            st.warning("Define al menos un eje de evaluación en la barra lateral.")
        else:
            prompt = build_interview_prompt(axes_list, scale_max)
            with st.spinner("Evaluando entrevista(s) con el modelo..."):
                scores_df, raw_response = handle_llm_errors(
                    call_llm_for_json, prompt, st.session_state.interview_text
                )
                if scores_df is not None:
                    st.session_state.scores_df = scores_df
                    st.session_state.raw_response_entrevista = raw_response
                    st.session_state.axes_list = axes_list
                    st.session_state.scale_max = scale_max

    if "scores_df" not in st.session_state:
        st.stop()

    st.divider()
    st.subheader("📋 Puntajes extraídos (editable)")
    st.caption("Corrige manualmente cualquier puntaje antes de generar el radar, si lo consideras necesario.")

    edited_scores = st.data_editor(
        st.session_state.scores_df, num_rows="dynamic", use_container_width=True, key="editor_scores"
    )

    with st.expander("Ver JSON crudo devuelto por el modelo"):
        st.code(st.session_state.get("raw_response_entrevista", ""), language="json")

    csv_buf2 = io.StringIO()
    edited_scores.to_csv(csv_buf2, index=False)
    st.download_button("⬇️ Descargar puntajes como CSV", csv_buf2.getvalue(),
                        file_name="puntajes_entrevista.csv", mime="text/csv")

    if edited_scores.empty:
        st.warning("La tabla de puntajes está vacía.")
        st.stop()

    # Detectar columna de identificación (Candidato) y columnas de ejes numéricos
    id_col = "Candidato" if "Candidato" in edited_scores.columns else edited_scores.columns[0]
    axis_cols = [c for c in st.session_state.axes_list if c in edited_scores.columns]
    if not axis_cols:
        axis_cols = [c for c in edited_scores.columns
                     if pd.to_numeric(edited_scores[c], errors="coerce").notna().all()]

    for c in axis_cols:
        edited_scores[c] = pd.to_numeric(edited_scores[c], errors="coerce")

    # --------------------------------------------------------------------------------
    # RADAR INTERACTIVO
    # --------------------------------------------------------------------------------
    st.divider()
    st.subheader("🕸️ Radar comparativo de candidatos")

    all_candidates = edited_scores[id_col].dropna().astype(str).tolist()
    selected = st.multiselect(
        "Candidatos a comparar", all_candidates, default=all_candidates, key="radar_select"
    )

    if selected and axis_cols:
        fig_radar = go.Figure()
        for _, row in edited_scores[edited_scores[id_col].astype(str).isin(selected)].iterrows():
            values = [row[c] for c in axis_cols]
            values_closed = values + values[:1]
            theta_closed = axis_cols + axis_cols[:1]
            fig_radar.add_trace(go.Scatterpolar(
                r=values_closed, theta=theta_closed, fill="toself",
                name=str(row[id_col]), opacity=0.6,
            ))
        fig_radar.update_layout(
            polar=dict(radialaxis=dict(visible=True, range=[0, st.session_state.scale_max])),
            showlegend=True, height=600,
            title="Comparación de competencias por candidato",
        )
        st.plotly_chart(fig_radar, use_container_width=True)
    else:
        st.info("Selecciona al menos un candidato y define ejes numéricos válidos para ver el radar.")

    st.markdown("**Tabla comparativa de puntajes**")
    st.dataframe(edited_scores, use_container_width=True)

    if len(all_candidates) >= 2 and axis_cols:
        st.markdown("**Promedio por eje (todos los candidatos evaluados)**")
        avg_scores = edited_scores[axis_cols].mean().reset_index()
        avg_scores.columns = ["Eje", "Promedio"]
        fig_avg = px.bar(avg_scores, x="Eje", y="Promedio", text="Promedio",
                          title="Promedio del grupo por eje de competencia")
        fig_avg.update_traces(texttemplate="%{text:.1f}")
        st.plotly_chart(fig_avg, use_container_width=True)
